"""Embeds chess positions to docx.

Setup:
  Install python version >= 3.7

  Requirements:
    pip install chess==1.9.3
    pip install cairosvg==2.5.2
    pip install python-docx==0.8.11
"""


__version__ = '0.4.2'


import random
import argparse

import chess
import chess.svg
import cairosvg
from docx import Document
from docx.shared import Inches


def get_epd(pos_file, is_shuffle):
    """Converts epd file to epd list.

    Args:
      fn (str): The file with epd positions.
      is_shuffle (bool): If true, will shuffle the positions.

    Returns:
      epds (list): A list of epd positions.
    """
    epds = []
    with open(pos_file, encoding='utf-8') as handle:
        for lines in handle:
            epd_line = lines.rstrip()
            epds.append(epd_line)

    if is_shuffle:
        random.shuffle(epds)

    return epds


def set_comments(document, board, epd_info,
                 show_fen, show_bm, show_id, show_c0):
    """Adds diagram comments"""
    parag = document.add_paragraph()

    if show_fen:
        run = parag.add_run(f'{board.fen()}')

    if show_bm:
        bms = epd_info.get('bm', None)
        if bms is not None:
            bms_sans = [board.san(m) for m in bms]
            bms_sans_str = ' '.join(bms_sans)
        else:
            bms_sans_str = None

        if show_fen:
            run.add_break()
        run = parag.add_run(f'bm: {bms_sans_str}')

    if show_id:
        if show_fen or show_bm:
            run.add_break()
        run = parag.add_run(f'id: {epd_info.get("id", None)}')

    if show_c0:
        if show_fen or show_bm or show_id:
            run.add_break()
        run = parag.add_run(f'c0: {epd_info.get("c0", None)}')

    if show_fen or show_bm or show_id or show_c0:
        run.add_break()  # vertical space gap


def get_board_orientation(board, board_orientation):
    """Returns board orientation"""
    if board_orientation == 'white':
        orient = chess.WHITE
    elif board_orientation == 'black':
        orient = chess.BLACK
    else:
        orient = board.turn

    return orient


def epd2doc(epd_file, output_file, max_pos, header,
            board_orientation, show_fen, show_bm,
            show_id, randomize_position, board_image_pixel_size,
            doc_image_inch_size, show_c0):
    """Embeds chess position diagram into the docx.

    Reads EPD or FEN file and embed it in docx. Adds comment
    on the diagram if flags are enabled.

    Args:
      epd_file (str): The file that contains chess position in EPD or FEN format.
      output_file (str): The docx file to write the output.

    Returns:
      None
    """
    epds = get_epd(epd_file, randomize_position)

    num_pos_printed = max(1, min(1e6, max_pos))
    pngfn = 'tmp_2Eqsm5b.png'
    board_size = board_image_pixel_size

    document = Document()
    document.add_heading(header, 0)

    for cnt, epd in enumerate(epds):
        board = chess.Board()
        epd_info = board.set_epd(epd)

        # Create svg object.
        mysvg = chess.svg.board(
            board,
            size=board_size,
            orientation=get_board_orientation(board, board_orientation))

        # Convert svg to png image file.
        cairosvg.svg2png(mysvg, write_to=pngfn)

        # Insert png image into docx.
        document.add_picture(pngfn, width=Inches(doc_image_inch_size))

        set_comments(document, board, epd_info, show_fen,
                     show_bm, show_id, show_c0)

        if cnt + 1 >= num_pos_printed:
            break

    try:
        document.save(output_file)
    except PermissionError:
        print('Error in saving the output. Please close the file if it is open.')


def main():
    parser = argparse.ArgumentParser()

    parser.add_argument('--epd-file', required=True, type=str,
                        help='The file name that contains epd or fen positions, (required=True).')
    parser.add_argument('--output-file', required=True, type=str,
                        help='The output docx filename, (required=True).')
    parser.add_argument('--board-orientation', required=False, type=str, default='side',
                        help='The diagram board orientation, (required=False, default=side, '
                             'values=[side, white, black]).')
    parser.add_argument('--max-pos', required=False, type=int, default=1000,
                        help='The maximum number of positions to embed in the doc, '
                             '(required=False, default=1000, min=1, max=1e6)')
    parser.add_argument('--board-image-pixel-size', required=False, default=None,
                        help='The size of the board image in pixel, '
                             '(required=False, default=None). e.g. --board-image-pixel-size 350')
    parser.add_argument('--doc-image-inch-size', required=False, type=float, default=3.0,
                        help='The size of the board image in inches in the doc, '
                             '(required=False, default=3.0).')
    parser.add_argument('--header', required=False, type=str, default='Chess Positions',
                        help='The Text that will appear at the top of the document, '
                             '(required=False, default="Chess Positions").')
    parser.add_argument('--show-fen', action='store_true',
                        help='A flag to show epd / fen in the doc.')
    parser.add_argument('--show-bm', action='store_true',
                        help='A flag to show bm in the doc.')
    parser.add_argument('--show-id', action='store_true',
                        help='A flag to show epd id in the doc.')
    parser.add_argument('--show-c0', action='store_true',
                        help='A flag to show epd c0 opcode value in the doc.')
    parser.add_argument('--randomize-position', action='store_true',
                        help='A flag to shuffle the positions before embedding to the doc.')
    parser.add_argument('-v', '--version', action='version', version=f'{__version__}')

    args = parser.parse_args()

    epd2doc(args.epd_file, args.output_file, args.max_pos, args.header,
            args.board_orientation, args.show_fen, args.show_bm,
            args.show_id, args.randomize_position, args.board_image_pixel_size,
            args.doc_image_inch_size, args.show_c0)


if __name__ == '__main__':
    main()
